#!/usr/bin/env python3
"""
Advanced PDF to DOCX Converter using Python libraries
This script provides superior conversion quality with layout preservation
"""

import sys
import json
import argparse
from pathlib import Path
import tempfile
import os

try:
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import PIL.Image as Image
    import io
    import base64
except ImportError as e:
    print(json.dumps({
        "error": f"Missing required Python library: {e}",
        "install_command": "pip install PyMuPDF python-docx Pillow"
    }))
    sys.exit(1)

class AdvancedPDFConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for better formatting"""
        styles = self.doc.styles
        
        # Title style
        if 'PDFTitle' not in [style.name for style in styles]:
            title_style = styles.add_style('PDFTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(18)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(0x2F, 0x5F, 0x8F)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        if 'PDFHeading' not in [style.name for style in styles]:
            heading_style = styles.add_style('PDFHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0x1F, 0x4F, 0x7F)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        
        # Body text style
        if 'PDFBody' not in [style.name for style in styles]:
            body_style = styles.add_style('PDFBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Calibri'
            body_style.font.size = Pt(11)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_style.paragraph_format.space_after = Pt(6)
    
    def analyze_text_block(self, block):
        """Analyze text block to determine its type and formatting"""
        text = block.get("text", "").strip()
        if not text:
            return None
        
        bbox = block.get("bbox", [0, 0, 0, 0])
        font_info = block.get("font", {})
        font_size = font_info.get("size", 11)
        font_flags = font_info.get("flags", 0)
        
        # Determine block type based on formatting
        is_bold = bool(font_flags & 2**4)  # Bold flag
        is_italic = bool(font_flags & 2**1)  # Italic flag
        
        block_type = "body"
        if font_size > 14 or is_bold:
            if len(text) < 100 and '\n' not in text:
                block_type = "heading"
        
        if text.isupper() and len(text) < 60:
            block_type = "title"
        
        return {
            "type": block_type,
            "text": text,
            "bold": is_bold,
            "italic": is_italic,
            "font_size": font_size,
            "bbox": bbox
        }
    
    def extract_images(self, pdf_doc, page_num):
        """Extract images from PDF page"""
        page = pdf_doc[page_num]
        image_list = page.get_images()
        images = []
        
        for img_index, img in enumerate(image_list):
            try:
                xref = img[0]
                pix = fitz.Pixmap(pdf_doc, xref)
                
                if pix.n - pix.alpha < 4:  # GRAY or RGB
                    img_data = pix.tobytes("png")
                    images.append({
                        "data": base64.b64encode(img_data).decode(),
                        "format": "png",
                        "bbox": img[1:5] if len(img) > 4 else None
                    })
                pix = None
            except Exception as e:
                print(f"Error extracting image {img_index}: {e}", file=sys.stderr)
                continue
        
        return images
    
    def convert_pdf_to_docx(self, pdf_path, output_path=None):
        """Convert PDF to DOCX with advanced formatting preservation"""
        try:
            pdf_doc = fitz.open(pdf_path)
            
            # Add document title
            title_para = self.doc.add_paragraph()
            title_para.style = 'PDFTitle'
            title_para.add_run(f"Converted from: {Path(pdf_path).stem}")
            
            # Add metadata
            meta_para = self.doc.add_paragraph()
            meta_para.add_run(f"Pages: {len(pdf_doc)} | Conversion Date: ")
            from datetime import datetime
            meta_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add separator
            self.doc.add_paragraph("─" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                
                # Extract text blocks with formatting
                blocks = page.get_text("dict")["blocks"]
                
                for block in blocks:
                    if "lines" in block:  # Text block
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            
                            if line_text.strip():
                                analyzed = self.analyze_text_block({
                                    "text": line_text,
                                    "bbox": line["bbox"],
                                    "font": line["spans"][0] if line["spans"] else {}
                                })
                                
                                if analyzed:
                                    self.add_formatted_text(analyzed)
                
                # Extract and add images
                images = self.extract_images(pdf_doc, page_num)
                for img in images:
                    self.add_image_to_doc(img)
                
                # Add page break (except for last page)
                if page_num < len(pdf_doc) - 1:
                    self.doc.add_page_break()
            
            pdf_doc.close()
            
            # Save document
            if not output_path:
                output_path = str(Path(pdf_path).with_suffix('.docx'))
            
            self.doc.save(output_path)
            
            return {
                "success": True,
                "output_path": output_path,
                "pages": len(pdf_doc),
                "size_kb": round(Path(output_path).stat().st_size / 1024, 2)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def add_formatted_text(self, analyzed_block):
        """Add text with appropriate formatting"""
        if analyzed_block["type"] == "title":
            para = self.doc.add_paragraph()
            para.style = 'PDFTitle'
            para.add_run(analyzed_block["text"])
        elif analyzed_block["type"] == "heading":
            para = self.doc.add_paragraph()
            para.style = 'PDFHeading'
            para.add_run(analyzed_block["text"])
        else:
            para = self.doc.add_paragraph()
            para.style = 'PDFBody'
            run = para.add_run(analyzed_block["text"])
            run.bold = analyzed_block["bold"]
            run.italic = analyzed_block["italic"]
    
    def add_image_to_doc(self, img_data):
        """Add image to document with proper sizing"""
        try:
            img_bytes = base64.b64decode(img_data["data"])
            img_stream = io.BytesIO(img_bytes)
            
            # Add image with reasonable size
            para = self.doc.add_paragraph()
            run = para.runs[0] if para.runs else para.add_run()
            
            # Calculate appropriate size (max 6 inches wide)
            try:
                with Image.open(img_stream) as pil_img:
                    width, height = pil_img.size
                    aspect_ratio = height / width
                    
                    max_width = Inches(6)
                    if width > 800:  # If image is large
                        doc_width = max_width
                        doc_height = max_width * aspect_ratio
                    else:
                        doc_width = Inches(width / 100)  # Scale down
                        doc_height = Inches(height / 100)
                    
                    img_stream.seek(0)
                    run.add_picture(img_stream, width=doc_width, height=doc_height)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
            except Exception as e:
                print(f"Error processing image: {e}", file=sys.stderr)
                
        except Exception as e:
            print(f"Error adding image to document: {e}", file=sys.stderr)

def main():
    parser = argparse.ArgumentParser(description='Advanced PDF to DOCX Converter')
    parser.add_argument('input_pdf', help='Input PDF file path')
    parser.add_argument('-o', '--output', help='Output DOCX file path')
    parser.add_argument('--json', action='store_true', help='Output JSON result')
    
    args = parser.parse_args()
    
    if not Path(args.input_pdf).exists():
        result = {"success": False, "error": "Input PDF file not found"}
    else:
        converter = AdvancedPDFConverter()
        result = converter.convert_pdf_to_docx(args.input_pdf, args.output)
    
    if args.json:
        print(json.dumps(result))
    else:
        if result["success"]:
            print(f"✅ Conversion successful: {result['output_path']}")
            print(f"📄 Pages: {result['pages']} | Size: {result['size_kb']} KB")
        else:
            print(f"❌ Conversion failed: {result['error']}")
            sys.exit(1)

if __name__ == "__main__":
    main()
